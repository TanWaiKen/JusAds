from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "JusAds_User_Manual.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
MUTED = "595959"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF4CC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(doc, text, *, size=11, color="000000", bold=False, after=7, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    set_font(p.add_run(text), size, color, bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(16 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13 if level == 2 else 12, BLUE if level < 3 else DARK, True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), 11)


def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(step), 11)


def add_note(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(f"{label}: "), 10, DARK, True)
    set_font(p.add_run(text), 10, "000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_feature_table(doc):
    rows = [
        ("Projects", "Organise generation and compliance work; share a project with a colleague."),
        ("Ad generation", "Create localised text, image, audio, or video advertising concepts."),
        ("Compliance review", "Review risks, explanation, rules, evidence, and suggested improvements."),
        ("Violation evidence", "Stores structured findings for audit summaries and timelines."),
        ("Remix", "Creates a remediation version and marks it pending recheck."),
        ("Content Ideas", "Shows cultural events and evidence-backed creative ideas."),
        ("Hook references", "Search YouTube Shorts references and learn preferred hook styles."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [2500, 6860]
    for i, title in enumerate(("Feature", "What it does")):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cell.paragraphs[0].add_run(title), 10, DARK, True)
    for feature, description in rows:
        cells = table.add_row().cells
        for i, text in enumerate((feature, description)):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_font(cells[i].paragraphs[0].add_run(text), 10, "000000", i == 0)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    for name, size, color in (("Normal", 11, "000000"), ("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("JusAds | User Manual"), 9, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("JusAds user guide - internal demonstration edition"), 9, MUTED)

    add_text(doc, "JusAds", size=28, color=DARK, bold=True, after=1)
    add_text(doc, "User Manual", size=20, color=BLUE, bold=True, after=12)
    add_text(doc, "Create local advertising, review compliance risks, and keep a clear evidence trail.", size=12, color=MUTED, after=18)
    add_note(doc, "Purpose", "Use this guide to operate JusAds safely. It explains what the system does today and labels capabilities that are still being completed.")
    add_text(doc, "Version: July 2026", size=10, color=MUTED, after=4)
    doc.add_page_break()

    add_heading(doc, "1. What JusAds does")
    add_text(doc, "JusAds helps Malaysian and regional SMEs create locally relevant ads, find creative inspiration, and check advertising material before publishing. It is a decision-support tool: users remain responsible for final approval and publication.")
    add_feature_table(doc)

    add_heading(doc, "2. Sign in and navigate")
    add_steps(doc, [
        "Open the JusAds web application and sign in with your authorised account.",
        "Use Home to view recent activity and project shortcuts.",
        "Use New Project to create a workspace for one campaign or client.",
        "Use Content Ideas to explore events, trends, and evidence-backed creative references.",
    ])
    add_note(doc, "Security", "Your account identity controls project access. Do not share your password or use another person's account.")

    add_heading(doc, "3. Manage projects and tasks")
    add_steps(doc, [
        "Select New Project and enter a clear campaign name.",
        "Open the project from the sidebar to view its execution history.",
        "Use New Task to start either ad generation or a compliance review.",
        "Use Share only when you are the project owner. Shared members can view or work according to their assigned access.",
    ])
    add_note(doc, "Good practice", "Keep one project per campaign or customer. This makes the audit trail, generated media, and compliance history easier to explain.")

    add_heading(doc, "4. Create an ad")
    add_steps(doc, [
        "Create a Generation task in the selected project.",
        "Add your product, target platform, target audience, and creative strategy.",
        "Choose a style and use the workspace to request copy, visuals, audio, or video concepts.",
        "Review generated output before downloading or using it externally.",
    ])
    add_note(doc, "Hook Reference", "The YouTube Shorts hook search is available for every creative style. Select a reference only when you have reviewed it and it suits your brand.")

    add_heading(doc, "5. Run a compliance review")
    add_steps(doc, [
        "Create a Compliance task and upload the ad material.",
        "Provide the market, audience, platform, and other requested context.",
        "Wait for the review result. The system presents risk, explanation, rules/evidence, and suggested changes.",
        "Open the task later from Execution History to review the saved result.",
    ])
    add_text(doc, "The system saves the complete result plus a normalised list of violations. The normalised entries support future timelines, reporting, and audit queries without changing the original AI result.")
    add_note(doc, "Important", "A compliance result is decision support, not legal advice or an automatic guarantee of regulatory approval.", LIGHT_GOLD)

    add_heading(doc, "6. Understand the compliance result")
    add_bullet(doc, "Risk level and percentage: a prioritisation signal, not a legal ruling.")
    add_bullet(doc, "High-risk indicators: actionable findings that should be reviewed before publishing.")
    add_bullet(doc, "Evidence and explanation: why the system raised the concern and what rule or context it used.")
    add_bullet(doc, "Suggestion/localisation plan: practical directions for improving the material.")
    add_bullet(doc, "Violation timeline: for supported media, where a concern occurs in the content.")

    add_heading(doc, "7. Remediate a finding")
    add_steps(doc, [
        "From a compliance result that needs changes, choose Remix/Remediate.",
        "Review the proposed revised material and the changes made.",
        "The remediated asset is saved as a distinct version, never overwriting the original.",
        "Its status becomes Pending Recheck until a new compliance evaluation completes.",
    ])
    add_note(doc, "Current limitation", "The durable recheck worker is still under development. Do not present a Pending Recheck asset as verified compliant. The required production flow is: pending job -> real evaluator -> immutable evaluation -> verified compliant or verified non-compliant.", LIGHT_GOLD)

    add_heading(doc, "8. Use Content Ideas and hook references")
    add_steps(doc, [
        "Open Content Ideas from the sidebar.",
        "Review upcoming cultural moments and current ad ideas.",
        "Open the linked evidence before using an idea in a campaign.",
        "Check the YouTube hook-reference message: it tells you whether company-context references are new, cached, or not yet available.",
        "Use a selected hook reference as inspiration; do not copy another creator's protected content.",
    ])

    add_heading(doc, "9. Troubleshooting")
    trouble = [
        ("No projects appear", "Refresh after signing in. If the problem continues, confirm the backend is running and contact the system administrator."),
        ("A feature says unavailable", "Try again later. The system keeps the original task and records a safe error rather than silently claiming success."),
        ("No YouTube references are cached", "This is normal before the first successful company-context lookup. It is a cache, not a missing user record."),
        ("Remediation is pending recheck", "Treat the revised asset as draft material until the real recheck completes."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [3000, 6360]
    for i, title in enumerate(("Situation", "What to do")):
        set_cell_width(table.cell(0, i), widths[i]); set_cell_shading(table.cell(0, i), LIGHT_BLUE); set_font(table.cell(0, i).paragraphs[0].add_run(title), 10, DARK, True)
    for left, right in trouble:
        cells = table.add_row().cells
        for i, text in enumerate((left, right)):
            set_cell_width(cells[i], widths[i]); set_font(cells[i].paragraphs[0].add_run(text), 10, "000000", i == 0)

    add_heading(doc, "10. Safe operating checklist")
    add_bullet(doc, "Confirm the correct project and target market before running generation or compliance checks.")
    add_bullet(doc, "Review source links and explanations; do not rely only on a colour or score.")
    add_bullet(doc, "Keep original and remediated assets separate for auditability.")
    add_bullet(doc, "Publish only after a human has approved the final material.")
    add_bullet(doc, "Do not claim verified remediation until the recheck status is verified compliant.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
